import os
import json
from typing import List, Dict, Optional

# Import from the new 'docx' library (python-docx-oss)
import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT


def configure_heading_styles(doc):
    """
    Configure the built-in Heading styles for base formatting.
    
    This preserves ToC functionality since Word's Table of Contents
    is based on these built-in Heading styles.
    """
    font_sizes = {
        'Heading 1': Pt(14),
        'Heading 2': Pt(13),
        'Heading 3': Pt(12),
        'Heading 4': Pt(11),
        'Heading 5': Pt(11),
        'Heading 6': Pt(11),
        'Heading 7': Pt(11),
        'Heading 8': Pt(11),
        'Heading 9': Pt(11),
    }
    
    for style_name, font_size in font_sizes.items():
        if style_name in doc.styles:
            style = doc.styles[style_name]
            
            style.font.bold = True
            style.font.underline = False
            style.font.size = font_size
            style.font.name = 'Calibri'
            style.font.color.rgb = None
            
            style.paragraph_format.space_before = Pt(12) if style_name == 'Heading 1' else Pt(8)
            style.paragraph_format.space_after = Pt(6)


def _estimate_column_widths_from_content(rows: list, num_cols: int) -> List[float]:
    """
    Estimate relative column widths based on the max content length in each column.
    Used when Excel column width metadata is unavailable (all defaults).
    
    Returns a list of relative widths (not inches — caller scales to page width).
    """
    max_lengths = [1] * num_cols  # minimum 1 to avoid zero-width
    
    for row_data in rows:
        for col_idx in range(min(len(row_data), num_cols)):
            cell_text = str(row_data[col_idx]) if row_data[col_idx] is not None else ""
            max_lengths[col_idx] = max(max_lengths[col_idx], len(cell_text))
    
    return max_lengths


# ---------------------------------------------------------------------------
# Shared table styling — single source of truth for all table formatting
# ---------------------------------------------------------------------------

TABLE_FONT_SIZE = Pt(8)
TABLE_FONT_NAME = 'Calibri'
HEADER_SHADING_COLOR = "D9E2F3"  # light blue-gray


def _set_cell_border(cell, **kwargs):
    """
    Set individual borders on a table cell.
    
    Usage: _set_cell_border(cell, top={"sz": 4, "val": "single", "color": "CCCCCC"})
    Pass val="none" to hide a border edge.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing borders element if present
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)

    tcBorders = OxmlElement('w:tcBorders')
    for edge_name, attrs in kwargs.items():
        edge = OxmlElement(f'w:{edge_name}')
        for attr_key, attr_val in attrs.items():
            edge.set(qn(f'w:{attr_key}'), str(attr_val))
        tcBorders.append(edge)
    tcPr.append(tcBorders)


def _apply_clean_borders(table, num_rows: int, num_cols: int):
    """
    Apply a clean borderless look: only thin horizontal rules between rows,
    with a slightly heavier rule under the header row.
    No vertical dividers, no outer box.
    """
    NONE = {"val": "none", "sz": "0", "color": "auto"}
    THIN = {"val": "single", "sz": "4", "color": "BFBFBF"}
    HEADER_BOTTOM = {"val": "single", "sz": "8", "color": "808080"}

    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            cell = table.cell(row_idx, col_idx)

            top = NONE
            if row_idx == 0:
                top = THIN  # light top edge on first row

            bottom = THIN
            if row_idx == 0:
                bottom = HEADER_BOTTOM  # heavier rule under header
            elif row_idx == num_rows - 1:
                bottom = THIN  # close off bottom

            _set_cell_border(
                cell,
                top=top,
                bottom=bottom,
                left=NONE,
                right=NONE,
            )


def _style_table_cell(cell, font_size: Pt = TABLE_FONT_SIZE, bold: bool = False):
    """
    Style a single cell's paragraph: set font, size, and remove extra spacing
    so rows stay compact and text doesn't word-wrap as aggressively.
    """
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        for run in paragraph.runs:
            run.font.size = font_size
            run.font.name = TABLE_FONT_NAME
            if bold:
                run.font.bold = True


def _shade_cell(cell, color: str):
    """Apply background shading to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)


def _set_table_width(table, total_width_inches: float):
    """Set the total table width in DXA, enable autofit layout, and clear table-level borders."""
    total_width_dxa = int(total_width_inches * 1440)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    # Set total width
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Enable autofit so Word can shrink columns to fit
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tblPr.append(layout)

    # Clear table-level borders (cell-level borders take precedence)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_excel_table_to_docx(doc, table_data: dict, total_width_inches: float = 6.5):
    """
    Adds a table to the docx document using data extracted from Excel.
    
    Formatting: compact 8pt text, no vertical grid lines, light horizontal rules,
    shaded header row. Total table width fills the page; Word auto-fits columns.
    
    Args:
        doc: python-docx Document instance
        table_data: {"columns": [{"width": float}, ...], "rows": [[val, ...], ...]}
        total_width_inches: Total table width (default 6.5" = US Letter with 1" margins)
    """
    rows = table_data.get("rows", [])
    columns = table_data.get("columns", [])
    
    if not rows:
        return
        
    num_rows = len(rows)
    num_cols = len(columns)
    
    # Default table style has no borders; we apply our own via _apply_clean_borders.
    table = doc.add_table(rows=num_rows, cols=num_cols)

    _set_table_width(table, total_width_inches)
    
    # Populate cells with compact styled text
    for row_idx, row_data in enumerate(rows):
        is_header = (row_idx == 0)
        for col_idx, cell_value in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell_text = str(cell_value) if cell_value is not None else ""
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_text
            _style_table_cell(cell, bold=is_header)
            if is_header:
                _shade_cell(cell, HEADER_SHADING_COLOR)

    _apply_clean_borders(table, num_rows, num_cols)

def add_isolated_landscape_table(doc, table_data: dict, caption_text: str = ""):
    """
    Isolates a table on a new landscape page, then reverts the document to portrait.
    Uses narrow margins (0.5") to maximize table space.
    Caption is rendered on the landscape page before the section break.
    """
    landscape_section = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape_section.orientation = WD_ORIENT.LANDSCAPE
    
    landscape_section.page_width, landscape_section.page_height = (
        landscape_section.page_height, 
        landscape_section.page_width
    )
    
    # Narrow margins for landscape tables: 0.5" all around
    landscape_section.left_margin = Inches(0.5)
    landscape_section.right_margin = Inches(0.5)
    landscape_section.top_margin = Inches(0.5)
    landscape_section.bottom_margin = Inches(0.5)
    
    # Content width: 11" page - 1" total margins = 10"
    landscape_content_width = 10.0
    
    # Determine which renderer to use based on the table_data format
    columns = table_data.get("columns", [])
    is_excel_format = columns and isinstance(columns[0], dict) and "width" in columns[0]
    
    if is_excel_format:
        add_excel_table_to_docx(doc, table_data, total_width_inches=landscape_content_width)
    else:
        landscape_dxa = int(landscape_content_width * 1440)
        from complex_table_schema import add_complex_table
        add_complex_table(doc, table_data, total_width_dxa=landscape_dxa)
    
    # Caption goes here — on the landscape page, before reverting to portrait
    if caption_text is not None and str(caption_text).strip():
        add_table_caption(doc, caption_text)
    
    # Revert to portrait with original margins
    portrait_section = doc.add_section(WD_SECTION.NEW_PAGE)
    portrait_section.orientation = WD_ORIENT.PORTRAIT
    
    portrait_section.page_width, portrait_section.page_height = (
        portrait_section.page_height, 
        portrait_section.page_width
    )
    portrait_section.left_margin = Inches(1.0)
    portrait_section.right_margin = Inches(1.0)
    portrait_section.top_margin = Inches(1.0)
    portrait_section.bottom_margin = Inches(1.0)


def add_section_heading(doc, section_number: str, topic: str, level: int = 1):
    """
    Adds a section heading using built-in heading styles for ToC.
    """
    heading_level = max(1, min(level, 9))
    style_name = f'Heading {heading_level}'
    
    # Safely fallback to Normal if the template lacks this heading style
    if style_name not in doc.styles:
        style_name = 'Normal'
        
    p = doc.add_paragraph(style=style_name)
    
    if section_number:
        run_number = p.add_run(str(section_number))
        run_number.bold = True
        run_number.underline = False
        
        if topic:
            run_space = p.add_run(" ")
            run_space.bold = True
            run_space.underline = False
            
    if topic:
        run_topic = p.add_run(str(topic))
        run_topic.bold = True
        run_topic.underline = True
    
    return p


def add_docx_table_from_data(doc, table_data: Dict):
    """
    Adds a table to the docx document from structured data.
    Delegates to the complex table schema renderer.
    """
    from complex_table_schema import add_docx_table_from_data as add_complex_table
    return add_complex_table(doc, table_data)


def add_figure_caption(doc, text: str):
    """Adds a true Word caption for a FIGURE."""
    style_name = 'Caption' if 'Caption' in doc.styles else 'Normal'
    p = doc.add_paragraph(style=style_name)
    p.add_run("Figure ")

    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'SEQ Figure \\* ARABIC'
    run._r.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

    p.add_run(f": {text}")


def add_table_caption(doc, text: str = ""):
    """
    Adds a true Word caption for a TABLE, which can be used for a Table of Tables.
    """
    p = doc.add_paragraph(style='Caption')
    p.add_run("Table ")

    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'SEQ Table \\* ARABIC'
    run._r.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

    if text and str(text).strip():
        p.add_run(f": {text}")


def add_equation_caption(doc, text: str = ""):
    """
    Adds a true Word caption for an EQUATION, which can be used for a Table of Equations.
    """
    p = doc.add_paragraph(style='Caption')
    p.add_run("Equation ")

    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'SEQ Equation \\* ARABIC'
    run._r.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

    if text and str(text).strip():
        p.add_run(f": {text}")


def add_bordered_paragraph(doc, text: str):
    """
    Adds a paragraph with a single-line border around it.
    """
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '4')
        border_el.set(qn('w:space'), '1')
        border_el.set(qn('w:color'), 'auto')
        pBdr.append(border_el)
    
    pPr.append(pBdr)


def add_title_page(doc, title_data: Dict):
    """
    Adds a formatted title page to the document using the extracted data.
    
    Layout order:
    1. Document title (centered, bold, large)
    2. Approval status with date (date in red if DRAFT)
    3. Distribution statement box (bordered)
    4. Export warning (in the box)
    5. Destruction notice (in the box)
    6. CONTROLLED BY entries
    7. CUI CATEGORY
    8. POC
    """
    from docx.shared import RGBColor
    
    # --- 1. Document Title ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.add_run('\n\n\n')  # Top spacing
    
    title_text = title_data.get('document_title', '')
    if title_text:
        run = p_title.add_run(title_text)
        run.bold = True
        run.font.size = Pt(16)
    
    doc.add_paragraph()  # Spacing after title
    
    # --- 2. Approval Status (with date potentially in red) ---
    approval_status = title_data.get('approval_status', '')
    approval_date = title_data.get('approval_date', '')
    
    if approval_status:
        p_approval = doc.add_paragraph()
        p_approval.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Check if we need to highlight the date in red
        is_draft = 'DRAFT' in approval_date.upper() if approval_date else False
        
        if approval_date and approval_date in approval_status:
            # Split the text around the date
            parts = approval_status.split(approval_date, 1)
            if len(parts) == 2:
                p_approval.add_run(parts[0])
                date_run = p_approval.add_run(approval_date)
                if is_draft:
                    date_run.font.color.rgb = RGBColor(255, 0, 0)
                p_approval.add_run(parts[1])
            else:
                p_approval.add_run(approval_status)
        else:
            p_approval.add_run(approval_status)
    
    doc.add_paragraph()  # Spacing
    
    # --- 3-5. Distribution Statement Box ---
    distribution = title_data.get('distribution_statement', '')
    export_warning = title_data.get('export_warning', '')
    destruction = title_data.get('destruction_notice', '')
    
    box_content = []
    if distribution:
        box_content.append(distribution)
    if export_warning:
        box_content.append(export_warning)
    if destruction:
        box_content.append(destruction)
    
    if box_content:
        for text in box_content:
            add_bordered_paragraph(doc, text)
    
    doc.add_paragraph()  # Spacing
    
    # --- 6. CONTROLLED BY entries ---
    controlled_by = title_data.get('controlled_by', [])
    if controlled_by:
        if isinstance(controlled_by, str):
            controlled_by = [controlled_by]
        for entry in controlled_by:
            p = doc.add_paragraph()
            run_label = p.add_run("CONTROLLED BY: ")
            run_label.bold = True
            p.add_run(str(entry))
    
    # --- 7. CUI Category ---
    cui_category = title_data.get('cui_category', '')
    if cui_category:
        p = doc.add_paragraph()
        run_label = p.add_run("CUI CATEGORY: ")
        run_label.bold = True
        p.add_run(cui_category)
    
    # --- 8. Point of Contact ---
    poc = title_data.get('point_of_contact', '')
    if poc:
        p = doc.add_paragraph()
        run_label = p.add_run("POC: ")
        run_label.bold = True
        p.add_run(poc)
    
    # --- Optional: Document date if present ---
    doc_date = title_data.get('document_date', '')
    if doc_date:
        p = doc.add_paragraph()
        run_label = p.add_run("DATE: ")
        run_label.bold = True
        p.add_run(doc_date)


def add_field(paragraph, field_text: str):
    """
    Adds a Word field (like PAGE or SEQ) to a paragraph.
    """
    run = paragraph.add_run()
    
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_text
    run._r.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)


def add_figure_caption(doc, text: str = ""):
    """
    Adds a true Word caption for a FIGURE, which can be used for a Table of Figures.
    """
    p = doc.add_paragraph(style='Caption')
    p.add_run("Figure ")

    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'SEQ Figure \\* ARABIC'
    run._r.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

    if text and str(text).strip():
        p.add_run(f": {text}")


def create_docx_from_elements(elements: List[Dict], output_filename: str, figures_image_folder: str, part_number: str, title_data: Optional[Dict]):
    """
    Creates a .docx file from document elements.
    Handles 'section', 'unassigned_text_block', 'figure', 'table', 'equation', and 'table_layout' element types.
    """
    doc = docx.Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    configure_heading_styles(doc)

    if title_data:
        add_title_page(doc, title_data)
        doc.add_page_break()

    # --- Add Header ---
    section = doc.sections[0]
    header = section.header
    
    for p in header.paragraphs:
        p_element = p._element
        p_element.getparent().remove(p_element)

    cui_p = header.add_paragraph()
    cui_run = cui_p.add_run('CUI')
    cui_run.bold = True
    cui_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_p = header.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_p.add_run(str(part_number))
    info_p.add_run(" | Page ")
    add_field(info_p, "PAGE")

    # --- Add Footer ---
    footer = section.footer
    for p in footer.paragraphs:
        p_element = p._element
        p_element.getparent().remove(p_element)

    footer_p = footer.add_paragraph()
    footer_run = footer_p.add_run('CUI')
    footer_run.bold = True
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for element in elements:
        element_type = str(element.get("type", "")).strip().lower()

        if element_type == "section":
            section_number = str(element.get("section_number", ""))
            topic = str(element.get("topic", ""))
            content = str(element.get("content", ""))
            
            level = len(section_number.split('.')) if section_number else 1
            heading_level = max(1, min(level, 9))
            
            if section_number or topic:
                add_section_heading(doc, section_number, topic, level=heading_level)
            
            if content:
                doc.add_paragraph(content)
        
        elif element_type == "unassigned_text_block":
            content = str(element.get("content", ""))
            if content:
                doc.add_paragraph(content)

        elif element_type == "figure":
            export_data = element.get("export") or {}
            image_filename = export_data.get("image_file")
            caption_text = str(element.get('caption_text', ''))
            
            if image_filename:
                image_path = os.path.join(figures_image_folder, image_filename)
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(6.0))
                    add_figure_caption(doc, caption_text)
                else:
                    p_error = doc.add_paragraph()
                    p_error.add_run(f"[Figure image not found: {image_filename}]").italic = True
                    add_figure_caption(doc, caption_text)
            else:
                p_error = doc.add_paragraph()
                p_error.add_run(f"[Figure has no image file]").italic = True
                add_figure_caption(doc, caption_text)
        
        elif element_type == "table":
            export_data = element.get("export") or {}
            image_filename = export_data.get("image_file")
            caption_text = str(element.get('caption_text', ''))
            table_data = element.get("table_data")
            render_as_image = element.get("_render_as_image", False)
            render_landscape = element.get("_render_landscape", False)

            is_valid_text_table = (
                not render_as_image and
                isinstance(table_data, dict) and 
                len(table_data.get("columns", [])) > 0 and 
                len(table_data.get("rows", [])) > 0
            )

            if is_valid_text_table:
                # Detect format: Excel-format has columns as [{"width": ...}],
                # complex_table_schema format has columns as [{"name": ...}] or ["str", ...]
                columns = table_data.get("columns", [])
                is_excel_format = columns and isinstance(columns[0], dict) and "width" in columns[0]

                if render_landscape:
                    # Caption is rendered inside the landscape section
                    add_isolated_landscape_table(doc, table_data, caption_text=caption_text)
                else:
                    if is_excel_format:
                        add_excel_table_to_docx(doc, table_data)
                    else:
                        add_docx_table_from_data(doc, table_data)
                    add_table_caption(doc, caption_text)
            elif image_filename:
                image_path = os.path.join(figures_image_folder, image_filename)
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(6.0))
                    add_table_caption(doc, caption_text)
                else:
                    p_error = doc.add_paragraph()
                    p_error.add_run(f"[Table image not found: {image_filename}]").italic = True
                    add_table_caption(doc, caption_text)
            else:
                p_placeholder = doc.add_paragraph()
                p_placeholder.add_run(f"[Table content not available]").italic = True
                add_table_caption(doc, caption_text)

        elif element_type == "equation":
            export_data = element.get("export") or {}
            image_filename = export_data.get("image_file")
            caption_text = str(element.get('caption_text', ''))
            
            if image_filename:
                image_path = os.path.join(figures_image_folder, image_filename)
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(4.0))
                    add_equation_caption(doc, caption_text)
                else:
                    p_error = doc.add_paragraph()
                    p_error.add_run(f"[Equation image not found: {image_filename}]").italic = True
                    add_equation_caption(doc, caption_text)
            else:
                p_placeholder = doc.add_paragraph()
                p_placeholder.add_run(f"[Equation content not available]").italic = True
                add_equation_caption(doc, caption_text)

        elif element_type in ("table_layout", "tab_layout"):
            export_data = element.get("export") or {}
            image_filename = export_data.get("image_file")
            table_data = element.get("table_data")
            render_as_image = element.get("_render_as_image", False)
            render_landscape = element.get("_render_landscape", False)

            is_valid_text_table = (
                not render_as_image and
                isinstance(table_data, dict) and 
                len(table_data.get("columns", [])) > 0 and 
                len(table_data.get("rows", [])) > 0
            )

            if is_valid_text_table:
                columns = table_data.get("columns", [])
                is_excel_format = columns and isinstance(columns[0], dict) and "width" in columns[0]

                if render_landscape:
                    # Pass caption_text=None to explicitly suppress captions for layout tables
                    add_isolated_landscape_table(doc, table_data, caption_text=None)
                else:
                    if is_excel_format:
                        add_excel_table_to_docx(doc, table_data)
                    else:
                        add_docx_table_from_data(doc, table_data)
                    # Notice we deliberately skip calling add_table_caption() here
            elif image_filename:
                image_path = os.path.join(figures_image_folder, image_filename)
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(6.0))
                else:
                    p_error = doc.add_paragraph()
                    p_error.add_run(f"[Layout table image not found: {image_filename}]").italic = True

    doc.save(output_filename)

def run_docx_creation(input_path: str, output_path: str, figures_base_path: str, doc_stem: str, title_data_path: str):
    """
    Loads final elements and title data to generate a DOCX file.
    
    Handles both old format (list of elements) and new format (dict with page_metadata and elements).
    """
    if not os.path.exists(input_path):
        print(f"Error: Input file not found at {input_path}. Skipping DOCX creation for {doc_stem}.")
        return

    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Handle both old and new data formats
    if isinstance(data, dict) and 'elements' in data:
        elements = data.get('elements', [])
        page_metadata = data.get('page_metadata', {})
        print(f"  - Loaded {len(elements)} elements (new format)")
        if page_metadata:
            print(f"  - Page metadata available for {len(page_metadata)} pages")
    else:
        elements = data if isinstance(data, list) else []
        print(f"  - Loaded {len(elements)} elements (legacy format)")
    
    title_data = None
    if os.path.exists(title_data_path):
        with open(title_data_path, 'r', encoding='utf-8') as f:
            title_data_list = json.load(f)
            if title_data_list:
                title_data = title_data_list[0]
    else:
        print(f"  - Warning: Title data file not found at {title_data_path}. Proceeding without a title page.")

    from asset_processor import resolve_asset_directory
    figure_image_folder = resolve_asset_directory(figures_base_path, doc_stem)
    if not figure_image_folder:
        print(f"  - [Warning] No image folder found for '{doc_stem}' in {figures_base_path}")
        figure_image_folder = os.path.join(figures_base_path, doc_stem)

    create_docx_from_elements(elements, output_path, figure_image_folder, doc_stem, title_data)
    print(f"  - DOCX document successfully created at {output_path}")


if __name__ == '__main__':
    doc_stem = ""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    input_file = os.path.join(script_dir, "..", "results", f"{doc_stem}_repaired.json")
    output_file = os.path.join(script_dir, "..", "results", f"{doc_stem}_final_oss.docx")
    figures_path = os.path.join(script_dir, "..", "iris_ocr", "CM_Spec_OCR_and_figtab_output", "exports")

    run_docx_creation(input_file, output_file, figures_path, doc_stem)