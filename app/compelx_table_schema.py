"""
complex_table_schema.py

Schema definition and Word renderer for complex/irregular tables.

Supports:
- Column and row spanning (merged cells)
- Per-cell alignment (horizontal and vertical)
- Per-cell shading/background colors
- Header row designation (repeated across pages)
- Column width control (absolute or proportional)
- Cell-level font overrides (bold, italic, size)
- Nested content (multiple paragraphs per cell)
- Row height hints
- Border style overrides per cell

Schema Design
=============
The schema uses an explicit grid model where every logical cell declares its
position and span. This avoids the ambiguity of "repeat this value for merged
rows" that the flat columns+rows format suffered from.

Example minimal table_data:
{
    "columns": [
        {"name": "Part No.", "width_pct": 25},
        {"name": "Description", "width_pct": 50},
        {"name": "QTY", "width_pct": 25}
    ],
    "rows": [
        {"cells": [
            {"text": "123-456"},
            {"text": "Screw, Pan Head"},
            {"text": "4"}
        ]}
    ]
}

Example complex table_data with merges:
{
    "columns": [
        {"name": "Category", "width_pct": 30},
        {"name": "Sub-A", "width_pct": 35},
        {"name": "Sub-B", "width_pct": 35}
    ],
    "header_rows": 1,
    "rows": [
        {
            "is_header": true,
            "cells": [
                {"text": "Category", "bold": true, "shading": "D9E2F3"},
                {"text": "Sub-A", "bold": true, "shading": "D9E2F3"},
                {"text": "Sub-B", "bold": true, "shading": "D9E2F3"}
            ]
        },
        {
            "cells": [
                {"text": "Mechanical", "rowspan": 2, "bold": true, "valign": "center"},
                {"text": "Torque specs"},
                {"text": "15 Nm"}
            ]
        },
        {
            "cells": [
                null,
                {"text": "Pressure rating"},
                {"text": "300 PSI"}
            ]
        },
        {
            "cells": [
                {"text": "Electrical", "colspan": 1, "bold": true},
                {"text": "Voltage & Current combined", "colspan": 2}
            ]
        }
    ]
}

Backward Compatibility
======================
The renderer accepts BOTH the old flat format:
    {"columns": ["A", "B"], "rows": [["x", "y"]]}
and the new rich format. It auto-detects which is in use.
"""

from dataclasses import dataclass, field
from typing import List, Dict, Optional, Union
from enum import Enum


# ---------------------------------------------------------------------------
# Schema types (for documentation / validation — the renderer works off dicts)
# ---------------------------------------------------------------------------

class HAlign(Enum):
    LEFT = "left"
    CENTER = "center"
    RIGHT = "right"


class VAlign(Enum):
    TOP = "top"
    CENTER = "center"
    BOTTOM = "bottom"


@dataclass
class CellSchema:
    """Schema for a single table cell."""
    # Content — either a simple string or a list of paragraph strings
    text: Union[str, List[str]] = ""

    # Spanning
    colspan: int = 1
    rowspan: int = 1

    # Formatting
    bold: bool = False
    italic: bool = False
    font_size_pt: Optional[float] = None  # None = inherit from doc default
    halign: str = "left"     # "left", "center", "right"
    valign: str = "top"      # "top", "center", "bottom"

    # Appearance
    shading: Optional[str] = None  # hex color e.g. "D9E2F3", None = no fill

    # None sentinel: a null cell means "this grid position is covered by a
    # spanning cell above or to the left — skip it."


@dataclass
class ColumnSchema:
    """Schema for a column definition."""
    name: str = ""
    width_pct: Optional[float] = None   # percentage of table width (0-100)
    width_dxa: Optional[int] = None     # absolute width in DXA (1440 = 1 inch)


@dataclass
class RowSchema:
    """Schema for a single row."""
    cells: List[Optional[CellSchema]] = field(default_factory=list)
    is_header: bool = False              # repeat this row on every page
    height_dxa: Optional[int] = None     # row height hint


@dataclass
class TableSchema:
    """Top-level table schema."""
    columns: List[ColumnSchema] = field(default_factory=list)
    rows: List[RowSchema] = field(default_factory=list)
    header_rows: int = 0                 # number of header rows to repeat
    total_width_dxa: int = 9360          # default = US Letter content width (1" margins)
    style: str = "Table Grid"


# ---------------------------------------------------------------------------
# Normalization: convert dict (from JSON) into validated native dicts
# ---------------------------------------------------------------------------

def is_legacy_format(table_data: Dict) -> bool:
    """Detect old flat format: columns is a list of strings, rows is list of lists."""
    columns = table_data.get("columns", [])
    rows = table_data.get("rows", [])
    if not columns:
        return False
    # Legacy: columns are plain strings and rows are plain lists
    return isinstance(columns[0], str) and rows and isinstance(rows[0], list)


def normalize_table_data(table_data: Dict) -> Dict:
    """
    Convert any supported table_data format into the canonical rich format.

    Accepts:
      - Legacy flat: {"columns": ["A","B"], "rows": [["x","y"]]}
      - New rich: {"columns": [{"name":"A", ...}], "rows": [{"cells": [...]}]}

    Returns the rich format dict (always).
    """
    if is_legacy_format(table_data):
        return _normalize_legacy(table_data)
    return table_data


def _normalize_legacy(table_data: Dict) -> Dict:
    """Upgrade legacy flat format to the rich schema."""
    col_names = table_data["columns"]
    raw_rows = table_data["rows"]

    columns = [{"name": name} for name in col_names]

    rows = []
    # First row as header
    header_cells = [{"text": name, "bold": True} for name in col_names]
    rows.append({"is_header": True, "cells": header_cells})

    for raw_row in raw_rows:
        cells = [{"text": str(v)} for v in raw_row]
        rows.append({"cells": cells})

    return {
        "columns": columns,
        "rows": rows,
        "header_rows": 1,
    }


# ---------------------------------------------------------------------------
# Renderer: dict → python-docx Table
# ---------------------------------------------------------------------------

def compute_column_widths_dxa(columns: List[Dict], total_width_dxa: int) -> List[int]:
    """
    Resolve column widths to absolute DXA values.

    Priority: width_dxa > width_pct > equal distribution.
    """
    n = len(columns)
    widths = [None] * n
    remaining = total_width_dxa
    unresolved = []

    for i, col in enumerate(columns):
        if col.get("width_dxa"):
            widths[i] = col["width_dxa"]
            remaining -= widths[i]
        elif col.get("width_pct"):
            w = int(total_width_dxa * col["width_pct"] / 100)
            widths[i] = w
            remaining -= w
        else:
            unresolved.append(i)

    if unresolved:
        per_col = max(remaining // len(unresolved), 360)  # minimum ~0.25 inch
        for i in unresolved:
            widths[i] = per_col

    return widths


def add_complex_table(doc, table_data: Dict, total_width_dxa: int = 9360):
    """
    Render a complex table into a python-docx Document.

    Args:
        doc: python-docx Document instance.
        table_data: Table data dict (legacy or rich format accepted).
        total_width_dxa: Total table width in DXA units (default: US Letter content area).

    Returns:
        The python-docx Table object that was added.
    """
    from docx.shared import Pt, Inches, Emu
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table_data = normalize_table_data(table_data)

    columns = table_data.get("columns", [])
    rows = table_data.get("rows", [])
    header_rows_count = table_data.get("header_rows", 0)
    style_name = table_data.get("style", "Table Grid")

    num_cols = len(columns)
    num_rows = len(rows)

    if num_cols == 0 or num_rows == 0:
        return None

    col_widths = compute_column_widths_dxa(columns, total_width_dxa)

    # Build the merge map: track which grid cells are "occupied" by a spanning cell
    # occupied[r][c] = (origin_r, origin_c) if covered by a span, else None
    occupied = [[None for _ in range(num_cols)] for _ in range(num_rows)]

    for r, row in enumerate(rows):
        c_logical = 0  # walks across physical columns
        for cell_data in row.get("cells", []):
            # Skip past any columns already occupied by a prior rowspan
            while c_logical < num_cols and occupied[r][c_logical] is not None:
                c_logical += 1
            if c_logical >= num_cols:
                break

            if cell_data is None:
                # Explicit null — skip (covered by a span declared elsewhere)
                c_logical += 1
                continue

            cs = cell_data.get("colspan", 1)
            rs = cell_data.get("rowspan", 1)

            # Mark all grid positions this cell covers
            for dr in range(rs):
                for dc in range(cs):
                    rr, cc = r + dr, c_logical + dc
                    if rr < num_rows and cc < num_cols:
                        occupied[rr][cc] = (r, c_logical)

            c_logical += cs

    # Create the table with the full grid
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = style_name
    table.autofit = False

    # Set table width
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_width_dxa))
    tblW.set(qn("w:type"), "dxa")
    # Remove existing tblW if present
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblPr.append(tblW)

    # Set column widths via tblGrid
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(tbl.index(tblPr) + 1, tblGrid)
    # Clear existing gridCol
    for gc in tblGrid.findall(qn("w:gridCol")):
        tblGrid.remove(gc)
    for w in col_widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)

    # -----------------------------------------------------------------------
    # Populate cells
    # -----------------------------------------------------------------------
    HALIGN_MAP = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    VALIGN_XML = {
        "top": "top",
        "center": "center",
        "bottom": "bottom",
    }

    for r, row_data in enumerate(rows):
        row_obj = table.rows[r]

        # Header row repeat
        is_header = row_data.get("is_header", False) or r < header_rows_count
        if is_header:
            trPr = row_obj._tr.get_or_add_trPr()
            tblHeader = OxmlElement("w:tblHeader")
            trPr.append(tblHeader)

        # Row height
        height_dxa = row_data.get("height_dxa")
        if height_dxa:
            trPr = row_obj._tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), str(height_dxa))
            trHeight.set(qn("w:hRule"), "atLeast")
            trPr.append(trHeight)

        c_logical = 0
        cell_iter = iter(row_data.get("cells", []))

        for c_physical in range(num_cols):
            cell_obj = row_obj.cells[c_physical]

            # Is this cell the top-left origin of a span, or is it covered?
            origin = occupied[r][c_physical]

            if origin is None:
                # No cell data mapped here — leave empty
                _set_cell_width(cell_obj, col_widths[c_physical])
                continue

            origin_r, origin_c = origin

            if (origin_r, origin_c) != (r, c_physical):
                # This cell is covered by a span from another cell.
                # python-docx merge handles the XML; we just need to set width.
                _set_cell_width(cell_obj, col_widths[c_physical])
                continue

            # This is an origin cell — find the matching cell_data
            # Walk the cell iterator, skipping nulls
            cell_data = None
            while True:
                try:
                    candidate = next(cell_iter)
                except StopIteration:
                    break
                if candidate is not None:
                    cell_data = candidate
                    break

            if cell_data is None:
                _set_cell_width(cell_obj, col_widths[c_physical])
                continue

            cs = cell_data.get("colspan", 1)
            rs = cell_data.get("rowspan", 1)

            # Perform merge if needed
            if cs > 1 or rs > 1:
                end_r = min(r + rs - 1, num_rows - 1)
                end_c = min(c_physical + cs - 1, num_cols - 1)
                merge_target = table.cell(end_r, end_c)
                cell_obj = cell_obj.merge(merge_target)

            # ---- Set cell width (sum of spanned columns) ----
            spanned_width = sum(col_widths[c_physical:c_physical + cs])
            _set_cell_width(cell_obj, spanned_width)

            # ---- Set cell content ----
            text = cell_data.get("text", "")
            paragraphs_text = text if isinstance(text, list) else [text]

            # Clear default empty paragraph
            for i, para_text in enumerate(paragraphs_text):
                if i == 0:
                    p = cell_obj.paragraphs[0]
                    p.clear()
                else:
                    p = cell_obj.add_paragraph()

                run = p.add_run(str(para_text))

                # Font overrides
                if cell_data.get("bold"):
                    run.bold = True
                if cell_data.get("italic"):
                    run.italic = True
                font_size = cell_data.get("font_size_pt")
                if font_size:
                    run.font.size = Pt(font_size)

                # Horizontal alignment
                halign = cell_data.get("halign", "left")
                p.alignment = HALIGN_MAP.get(halign, WD_ALIGN_PARAGRAPH.LEFT)

            # ---- Vertical alignment ----
            valign = cell_data.get("valign", "top")
            tc = cell_obj._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign_el = OxmlElement("w:vAlign")
            vAlign_el.set(qn("w:val"), VALIGN_XML.get(valign, "top"))
            # Remove existing vAlign
            for existing in tcPr.findall(qn("w:vAlign")):
                tcPr.remove(existing)
            tcPr.append(vAlign_el)

            # ---- Shading ----
            shading_color = cell_data.get("shading")
            if shading_color:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), shading_color)
                # Remove existing shading
                for existing in tcPr.findall(qn("w:shd")):
                    tcPr.remove(existing)
                tcPr.append(shd)

    return table


def _set_cell_width(cell, width_dxa: int):
    """Set the preferred width on a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    # Remove existing
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcPr.append(tcW)


# ---------------------------------------------------------------------------
# Convenience: drop-in replacement for the old add_docx_table_from_data
# ---------------------------------------------------------------------------

def add_docx_table_from_data(doc, table_data: Dict):
    """
    Drop-in replacement for docx_writer.add_docx_table_from_data.
    Accepts both legacy flat format and the new rich format.
    """
    return add_complex_table(doc, table_data)