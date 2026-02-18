"""
test_complex_tables.py

Generates a .docx file with several complex tables to verify the schema
and renderer work correctly. Open the output in Word to inspect.
"""

import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

import docx
from docx.shared import Pt
from complex_table_schema import add_complex_table, add_docx_table_from_data


def make_test_doc():
    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ------------------------------------------------------------------
    # TEST 1: Legacy format backward compatibility
    # ------------------------------------------------------------------
    doc.add_paragraph("TEST 1: Legacy Flat Format (backward compat)", style="Heading 1")
    legacy_data = {
        "columns": ["Part No.", "Description", "QTY"],
        "rows": [
            ["123-456", "Screw, Pan Head", "4"],
            ["789-012", "Washer, Flat", "4"],
            ["345-678", "Nut, Hex", "8"],
        ],
    }
    add_docx_table_from_data(doc, legacy_data)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # TEST 2: Column span (colspan)
    # ------------------------------------------------------------------
    doc.add_paragraph("TEST 2: Column Spanning", style="Heading 1")
    colspan_data = {
        "columns": [
            {"name": "A", "width_pct": 25},
            {"name": "B", "width_pct": 25},
            {"name": "C", "width_pct": 25},
            {"name": "D", "width_pct": 25},
        ],
        "header_rows": 1,
        "rows": [
            {
                "is_header": True,
                "cells": [
                    {"text": "Col A", "bold": True, "shading": "D9E2F3"},
                    {"text": "Col B", "bold": True, "shading": "D9E2F3"},
                    {"text": "Col C", "bold": True, "shading": "D9E2F3"},
                    {"text": "Col D", "bold": True, "shading": "D9E2F3"},
                ],
            },
            {
                "cells": [
                    {"text": "Spans 2 columns", "colspan": 2, "shading": "FFF2CC"},
                    {"text": "C1"},
                    {"text": "D1"},
                ],
            },
            {
                "cells": [
                    {"text": "A2"},
                    {"text": "Spans 3 columns", "colspan": 3, "halign": "center", "shading": "E2EFDA"},
                ],
            },
            {
                "cells": [
                    {"text": "Full width row", "colspan": 4, "bold": True, "halign": "center", "shading": "FCE4D6"},
                ],
            },
        ],
    }
    add_complex_table(doc, colspan_data)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # TEST 3: Row span (rowspan)
    # ------------------------------------------------------------------
    doc.add_paragraph("TEST 3: Row Spanning", style="Heading 1")
    rowspan_data = {
        "columns": [
            {"name": "Category", "width_pct": 30},
            {"name": "Item", "width_pct": 40},
            {"name": "Value", "width_pct": 30},
        ],
        "header_rows": 1,
        "rows": [
            {
                "is_header": True,
                "cells": [
                    {"text": "Category", "bold": True, "shading": "D9E2F3"},
                    {"text": "Item", "bold": True, "shading": "D9E2F3"},
                    {"text": "Value", "bold": True, "shading": "D9E2F3"},
                ],
            },
            {
                "cells": [
                    {"text": "Mechanical", "rowspan": 3, "bold": True, "valign": "center", "shading": "E2EFDA"},
                    {"text": "Torque"},
                    {"text": "15 Nm"},
                ],
            },
            {
                "cells": [
                    None,  # covered by rowspan above
                    {"text": "Pressure"},
                    {"text": "300 PSI"},
                ],
            },
            {
                "cells": [
                    None,  # covered by rowspan above
                    {"text": "Temperature"},
                    {"text": "250°F"},
                ],
            },
            {
                "cells": [
                    {"text": "Electrical", "rowspan": 2, "bold": True, "valign": "center", "shading": "FCE4D6"},
                    {"text": "Voltage"},
                    {"text": "28V DC"},
                ],
            },
            {
                "cells": [
                    None,
                    {"text": "Current"},
                    {"text": "5A max"},
                ],
            },
        ],
    }
    add_complex_table(doc, rowspan_data)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # TEST 4: Mixed rowspan + colspan
    # ------------------------------------------------------------------
    doc.add_paragraph("TEST 4: Mixed Row + Column Spanning", style="Heading 1")
    mixed_data = {
        "columns": [
            {"name": "A", "width_pct": 20},
            {"name": "B", "width_pct": 20},
            {"name": "C", "width_pct": 20},
            {"name": "D", "width_pct": 20},
            {"name": "E", "width_pct": 20},
        ],
        "rows": [
            {
                "is_header": True,
                "cells": [
                    {"text": "Header spanning all 5", "colspan": 5, "bold": True,
                     "halign": "center", "shading": "2F5496", "font_size_pt": 13},
                ],
            },
            {
                "cells": [
                    {"text": "Big block\n(2x2)", "rowspan": 2, "colspan": 2,
                     "valign": "center", "halign": "center", "shading": "D9E2F3"},
                    {"text": "C1"},
                    {"text": "D1"},
                    {"text": "E1"},
                ],
            },
            {
                "cells": [
                    None, None,  # covered by 2x2 block
                    {"text": "C2"},
                    {"text": "Wide D+E", "colspan": 2, "halign": "center", "shading": "FFF2CC"},
                ],
            },
            {
                "cells": [
                    {"text": "A3"},
                    {"text": "B3"},
                    {"text": "Tall C\n(3 rows)", "rowspan": 3, "valign": "center", "shading": "E2EFDA"},
                    {"text": "D3"},
                    {"text": "E3"},
                ],
            },
            {
                "cells": [
                    {"text": "A4"},
                    {"text": "B4"},
                    None,  # covered by rowspan
                    {"text": "D4"},
                    {"text": "E4"},
                ],
            },
            {
                "cells": [
                    {"text": "Footer spans A+B", "colspan": 2, "bold": True, "shading": "FCE4D6"},
                    None,  # covered by rowspan
                    {"text": "Footer D+E", "colspan": 2, "bold": True, "shading": "FCE4D6"},
                ],
            },
        ],
    }
    add_complex_table(doc, mixed_data)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # TEST 5: Alignment and formatting showcase
    # ------------------------------------------------------------------
    doc.add_paragraph("TEST 5: Alignment & Formatting", style="Heading 1")
    fmt_data = {
        "columns": [
            {"name": "Left", "width_pct": 33},
            {"name": "Center", "width_pct": 34},
            {"name": "Right", "width_pct": 33},
        ],
        "rows": [
            {
                "is_header": True,
                "cells": [
                    {"text": "Left-aligned", "bold": True, "halign": "left", "shading": "D9E2F3"},
                    {"text": "Center-aligned", "bold": True, "halign": "center", "shading": "D9E2F3"},
                    {"text": "Right-aligned", "bold": True, "halign": "right", "shading": "D9E2F3"},
                ],
            },
            {
                "height_dxa": 800,
                "cells": [
                    {"text": "Top-left", "valign": "top", "halign": "left"},
                    {"text": "Center-center", "valign": "center", "halign": "center", "bold": True},
                    {"text": "Bottom-right", "valign": "bottom", "halign": "right", "italic": True},
                ],
            },
            {
                "cells": [
                    {"text": "Normal 11pt"},
                    {"text": "Small 9pt", "font_size_pt": 9},
                    {"text": "Large 14pt", "font_size_pt": 14, "bold": True},
                ],
            },
        ],
    }
    add_complex_table(doc, fmt_data)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # TEST 6: Multi-paragraph cell content
    # ------------------------------------------------------------------
    doc.add_paragraph("TEST 6: Multi-Paragraph Cells", style="Heading 1")
    multi_para_data = {
        "columns": [
            {"name": "Requirement", "width_pct": 30},
            {"name": "Description", "width_pct": 70},
        ],
        "rows": [
            {
                "is_header": True,
                "cells": [
                    {"text": "Requirement", "bold": True, "shading": "D9E2F3"},
                    {"text": "Description", "bold": True, "shading": "D9E2F3"},
                ],
            },
            {
                "cells": [
                    {"text": "REQ-001", "bold": True, "valign": "top"},
                    {"text": [
                        "The system shall support concurrent users.",
                        "Maximum latency shall not exceed 200ms under load.",
                        "See Section 4.2 for performance benchmarks.",
                    ]},
                ],
            },
            {
                "cells": [
                    {"text": "REQ-002", "bold": True, "valign": "top"},
                    {"text": [
                        "All data at rest shall be encrypted using AES-256.",
                        "Key rotation shall occur every 90 days.",
                    ]},
                ],
            },
        ],
    }
    add_complex_table(doc, multi_para_data)

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------
    output_path = "/home/claude/complex_table_tests.docx"
    doc.save(output_path)
    print(f"Test document saved to {output_path}")
    return output_path


if __name__ == "__main__":
    make_test_doc()