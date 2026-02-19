"""
test_ocr_to_table.py

Loads OCR JSON, converts through ocr_to_table, and renders via
complex_table_maker to verify the full pipeline.

Lives in: root/tests/
Run from project root:
    python -m tests.test_ocr_to_table

Expects test data at: root/tests/test_data/table_json_example.json
"""

import os
import sys
import json
from typing import Dict

# Add app directory to path so we can import modules directly
_TESTS_DIR = os.path.dirname(os.path.abspath(__file__))
_ROOT_DIR = os.path.dirname(_TESTS_DIR)
sys.path.insert(0, os.path.join(_ROOT_DIR, "app"))

import docx
from docx.shared import Pt

from ocr_to_table import convert_ocr_to_table_schema, convert_and_strip_empty
from consolidate_columns import consolidate_ocr_data
from complex_table_maker import add_complex_table


OUTPUT_DIR = os.path.join(_TESTS_DIR, "test_outputs")
TEST_DATA_DIR = os.path.join(_TESTS_DIR, "test_data")


def test_full_grid(ocr_data: Dict, doc):
    """Render the full grid as-is (no stripping, no consolidation)."""
    doc.add_paragraph("Full Grid (no trimming)", style="Heading 1")
    doc.add_paragraph(
        f"Grid: {ocr_data['n_rows']}x{ocr_data['n_cols']} | "
        f"Layout: {ocr_data.get('layout_type', '?')} | "
        f"Confidence: {ocr_data.get('confidence', 0):.1%}"
    )
    schema = convert_ocr_to_table_schema(ocr_data)
    add_complex_table(doc, schema)
    doc.add_paragraph()


def test_consolidated_grid(ocr_data: Dict, doc):
    """Render with column consolidation + edge stripping."""
    doc.add_paragraph("Consolidated Grid (phantom columns merged)", style="Heading 1")

    consolidated = consolidate_ocr_data(ocr_data)
    doc.add_paragraph(
        f"Original: {ocr_data['n_rows']}x{ocr_data['n_cols']} → "
        f"Consolidated: {consolidated['n_rows']}x{consolidated['n_cols']} cols"
    )

    groups = consolidated.get("_column_groups", [])
    for i, group in enumerate(groups):
        doc.add_paragraph(f"  Logical col {i}: grid cols {group}", style="List Bullet")

    schema = convert_and_strip_empty(ocr_data, consolidate=True)
    n_cols = len(schema["columns"])
    n_rows = len(schema["rows"])
    doc.add_paragraph(f"After strip: {n_rows} rows x {n_cols} cols")
    add_complex_table(doc, schema)
    doc.add_paragraph()

    return schema


def test_stripped_grid(ocr_data: Dict, doc):
    """Render with edge stripping only (no consolidation) for comparison."""
    doc.add_paragraph("Stripped Grid (edges only, no consolidation)", style="Heading 1")
    schema = convert_and_strip_empty(ocr_data, consolidate=False)
    n_cols = len(schema["columns"])
    n_rows = len(schema["rows"])
    doc.add_paragraph(f"Trimmed to: {n_rows} rows x {n_cols} cols")
    add_complex_table(doc, schema)
    doc.add_paragraph()


def test_rowspan(doc):
    """
    Regression test for the rowspan cell-loss bug.

    The old iterator-based renderer would desync when a row had
    [None, cell, cell, cell] with col 0 occupied by a rowspan.
    The None caused a double-skip, shifting everything right and
    dropping the last cell.
    """
    doc.add_paragraph("Rowspan Regression Test", style="Heading 1")
    doc.add_paragraph(
        "Verify all 4 columns render in row 2. "
        "Old bug: 'First item' was lost."
    )
    schema = {
        "columns": [
            {"name": "Cat", "width_pct": 20},
            {"name": "Name", "width_pct": 30},
            {"name": "Val", "width_pct": 20},
            {"name": "Notes", "width_pct": 30},
        ],
        "rows": [
            {
                "is_header": True,
                "cells": [
                    {"text": "Category", "rowspan": 2, "bold": True, "shading": "D9E2F3"},
                    {"text": "Name", "bold": True, "shading": "D9E2F3"},
                    {"text": "Value", "bold": True, "shading": "D9E2F3"},
                    {"text": "Notes", "bold": True, "shading": "D9E2F3"},
                ],
            },
            {
                "cells": [
                    None,
                    {"text": "Item A"},
                    {"text": "100"},
                    {"text": "First item"},
                ],
            },
            {
                "cells": [
                    {"text": "Electrical", "rowspan": 2, "bold": True},
                    {"text": "Voltage"},
                    {"text": "120V"},
                    {"text": "Standard"},
                ],
            },
            {
                "cells": [
                    None,
                    {"text": "Current"},
                    {"text": "15A"},
                    {"text": "Max rated"},
                ],
            },
        ],
        "header_rows": 1,
    }
    add_complex_table(doc, schema)
    doc.add_paragraph()


def test_schema_dump(ocr_data: Dict):
    """Print converted schema summary for inspection."""
    schema = convert_and_strip_empty(ocr_data, consolidate=True)

    total_cells = 0
    spanned_cells = 0
    low_conf = 0
    for row in schema["rows"]:
        for cell in row.get("cells", []):
            if cell is None:
                continue
            total_cells += 1
            if cell.get("colspan", 1) > 1 or cell.get("rowspan", 1) > 1:
                spanned_cells += 1
            if cell.get("italic"):
                low_conf += 1

    print(f"  Columns: {len(schema['columns'])}")
    print(f"  Rows: {len(schema['rows'])}")
    print(f"  Cells with content: {total_cells}")
    print(f"  Spanned cells: {spanned_cells}")
    print(f"  Low confidence (italic): {low_conf}")
    print(f"  Header rows: {schema['header_rows']}")
    print(f"  Col widths: {[c.get('width_pct') for c in schema['columns']]}")

    return schema


def main():
    json_path = os.path.join(TEST_DATA_DIR, "table_json_example.json")

    with open(json_path, "r", encoding="utf-8") as f:
        ocr_data = json.load(f)

    print(f"Loaded: {json_path}")
    print(f"Source grid: {ocr_data['n_rows']}x{ocr_data['n_cols']}")
    print()

    # Schema inspection
    print("Schema summary (consolidated + stripped):")
    schema = test_schema_dump(ocr_data)
    print()

    # Build docx
    doc = docx.Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_paragraph("OCR to Table Schema Test", style="Title")
    doc.add_paragraph(f"Source: {ocr_data.get('source_image', '?')}")
    doc.add_paragraph()

    test_full_grid(ocr_data, doc)
    consolidated_schema = test_consolidated_grid(ocr_data, doc)
    test_stripped_grid(ocr_data, doc)
    test_rowspan(doc)

    # Save outputs
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    docx_path = os.path.join(OUTPUT_DIR, "ocr_table_test.docx")
    doc.save(docx_path)
    print(f"Docx: {docx_path}")

    schema_path = os.path.join(OUTPUT_DIR, "ocr_table_schema_output.json")
    with open(schema_path, "w", encoding="utf-8") as f:
        json.dump(consolidated_schema, f, indent=2)
    print(f"Schema JSON: {schema_path}")


if __name__ == "__main__":
    main()